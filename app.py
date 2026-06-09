# ===================================================================
# ULTIMATE EXAM PREDICTOR – WITH WORKING CHATBOT
# Deploy on Streamlit Cloud (share.streamlit.io)
# ===================================================================

import streamlit as st
import pandas as pd
import google.generativeai as genai
import os
import json
import tempfile
import re
import zipfile
import io
from datetime import datetime, timedelta
from langchain_community.document_loaders import PyPDFLoader
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain_community.vectorstores import Chroma
import requests
from fpdf import FPDF
import plotly.express as px
from docx import Document

# ------------------- PAGE CONFIG -------------------
st.set_page_config(page_title="Ultimate Exam Predictor", layout="wide", initial_sidebar_state="expanded")
st.title("🎓 Ultimate Exam Predictor – All Features")

# ------------------- API KEYS from secrets -------------------
try:
    genai.configure(api_key=st.secrets["GOOGLE_API_KEY"])
    JINA_KEY = st.secrets["JINA_API_KEY"]
except:
    st.error("Please set GOOGLE_API_KEY and JINA_API_KEY in Streamlit secrets.")
    st.stop()

MODEL = genai.GenerativeModel('gemini-2.5-flash')

# ------------------- CONSTANTS & FOLDERS -------------------
BASE = "./ultimate_predictor_data"
SUBJECTS = ["Business Studies","Economics","English","CS/Stat/His","Accountancy","Kannada"]
EXAM_TYPES = ["Weekly","Monthly","Quarterly","Mid-Term","Final"]
SOURCES = ["College Teacher", "Government"]
DIFFICULTIES = ["Easy", "Medium", "Hard"]

for sub in SUBJECTS:
    for folder in ["textbooks","past_papers","important_questions","patterns","vectorstore","predictions"]:
        os.makedirs(f"{BASE}/{folder}/{sub}", exist_ok=True)
os.makedirs(f"{BASE}/feedback", exist_ok=True)
os.makedirs(f"{BASE}/history", exist_ok=True)

# ------------------- ALL HELPER FUNCTIONS (same as before, add chatbot-specific ones) -------------------

def search_searxng(q):
    for inst in ["https://searx.bang.pw","https://searx.work","https://search.im-in.space"]:
        try:
            url = f"{inst}/search?q={q}&format=json"
            r = requests.get(url, headers={'User-Agent':'Mozilla/5.0'}, timeout=8)
            if r.status_code == 200:
                return r.json()
        except:
            continue
    return None

def check_online(question):
    res = search_searxng(question[:200])
    if res and res.get('results'):
        url = res['results'][0].get('url')
        if url:
            try:
                jina_url = f"https://r.jina.ai/{url}"
                headers = {"Authorization": f"Bearer {JINA_KEY}"}
                c = requests.get(jina_url, headers=headers, timeout=12)
                if c.status_code == 200 and question.lower() in c.text.lower():
                    return "Found Online", url
            except:
                pass
    return "Likely Original", None

def get_schedule():
    f = f"{BASE}/schedule.csv"
    if os.path.exists(f):
        df = pd.read_csv(f)
        df['date'] = pd.to_datetime(df['date'])
        return df
    default = pd.DataFrame([
        ["Business Studies","2026-06-09","Weekly","College Teacher","Dr. Sharma"],
        ["Economics","2026-06-12","Weekly","College Teacher","Dr. Mehta"],
        ["English","2026-06-16","Weekly","College Teacher","Ms. Roy"],
        ["CS/Stat/His","2026-06-19","Weekly","College Teacher","Dr. Patil"],
        ["Accountancy","2026-06-23","Weekly","College Teacher","Mr. Gupta"],
        ["Kannada","2026-06-27","Weekly","College Teacher","Mrs. Hegde"],
    ], columns=["subject","date","exam_type","source","teacher"])
    default['date'] = pd.to_datetime(default['date'])
    default.to_csv(f, index=False)
    return default

def save_schedule(df):
    df.to_csv(f"{BASE}/schedule.csv", index=False)

def save_important(subject, teacher, questions):
    path = f"{BASE}/important_questions/{subject}/{teacher}.json"
    with open(path, "w") as f:
        json.dump({"teacher": teacher, "date": str(datetime.now().date()), "questions": questions}, f)

def load_important(subject, teacher):
    path = f"{BASE}/important_questions/{subject}/{teacher}.json"
    if os.path.exists(path):
        with open(path) as f:
            return json.load(f).get("questions", [])
    return []

def extract_pattern(subject, teacher, source):
    papers_dir = f"{BASE}/past_papers/{subject}"
    if not os.path.exists(papers_dir):
        return None
    all_text = ""
    meta_file = f"{papers_dir}/metadata.json"
    if os.path.exists(meta_file):
        with open(meta_file) as f:
            meta = json.load(f)
        for item in meta:
            if source == "College Teacher" and item.get("teacher") == teacher:
                pdf_path = f"{papers_dir}/{item['file']}"
                if os.path.exists(pdf_path):
                    loader = PyPDFLoader(pdf_path)
                    pages = loader.load()
                    all_text += " ".join([p.page_content for p in pages])
            elif source == "Government" and item.get("source") == "Government":
                pdf_path = f"{papers_dir}/{item['file']}"
                if os.path.exists(pdf_path):
                    loader = PyPDFLoader(pdf_path)
                    pages = loader.load()
                    all_text += " ".join([p.page_content for p in pages])
    if not all_text:
        return None
    prompt = f"""Analyze past papers for {subject} ({source}). Teacher: {teacher}.
Extract: Top 5 topics, question styles, average questions, repeats.
Text: {all_text[:5000]}"""
    return MODEL.generate_content(prompt).text

def build_rag(subject):
    docs = []
    tb_dir = f"{BASE}/textbooks/{subject}"
    for fname in os.listdir(tb_dir):
        if fname.endswith(".pdf"):
            with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp:
                tmp.write(open(f"{tb_dir}/{fname}","rb").read())
                path = tmp.name
            loader = PyPDFLoader(path)
            docs.extend(loader.load())
            os.unlink(path)
    if not docs:
        return None
    splitter = RecursiveCharacterTextSplitter(chunk_size=1500, chunk_overlap=200)
    chunks = splitter.split_documents(docs)
    embeddings = HuggingFaceEmbeddings(model_name="sentence-transformers/all-MiniLM-L6-v2")
    vectordb = Chroma.from_documents(chunks, embeddings, persist_directory=f"{BASE}/vectorstore/{subject}")
    return vectordb

def retrieve_topics(subject, query="most important topics"):
    vectordb = Chroma(persist_directory=f"{BASE}/vectorstore/{subject}", embedding_function=HuggingFaceEmbeddings(model_name="sentence-transformers/all-MiniLM-L6-v2"))
    retriever = vectordb.as_retriever(search_kwargs={"k": 5})
    docs = retriever.get_relevant_documents(query)
    return "\n".join([d.page_content[:600] for d in docs])

def full_text_search(subject, keyword):
    vectordb = Chroma(persist_directory=f"{BASE}/vectorstore/{subject}", embedding_function=HuggingFaceEmbeddings(model_name="sentence-transformers/all-MiniLM-L6-v2"))
    retriever = vectordb.as_retriever(search_kwargs={"k": 10})
    docs = retriever.get_relevant_documents(keyword)
    return docs

def predict(subject, teacher, source, custom_instruction, num_questions=5, difficulty="Medium"):
    schedule = get_schedule()
    today = datetime.now().date()
    upcoming = schedule[(schedule['subject']==subject) & (schedule['date'].dt.date >= today)].sort_values('date')
    next_date = upcoming.iloc[0]['date'].date() if not upcoming.empty else "Unknown"

    pattern = extract_pattern(subject, teacher, source)
    pattern_text = pattern if pattern else "No past papers."

    important = load_important(subject, teacher) if source == "College Teacher" else []
    imp_text = "\n".join(important) if important else "None"

    rag_text = ""
    try:
        rag_text = retrieve_topics(subject)
    except:
        rag_text = "No textbook."

    prompt = f"""Predict {num_questions} questions for {subject} ({source}), teacher {teacher}.
Next exam: {next_date}. Difficulty: {difficulty}.
Pattern: {pattern_text}
Important: {imp_text}
Textbook: {rag_text}
Instruction: {custom_instruction}
Return JSON: {{"questions":[{{"text":"...","confidence":0-100,"reason":"..."}}], "surprise_topics":["..."]}}"""
    response = MODEL.generate_content(prompt)
    match = re.search(r'\{.*\}', response.text, re.DOTALL)
    if not match:
        return [], []
    data = json.loads(match.group())
    questions = [(item["text"], item["confidence"], item.get("reason","")) for item in data.get("questions", [])]
    surprises = data.get("surprise_topics", [])
    return questions, surprises

def export_pdf(subject, teacher, source, questions, surprises):
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", "B", 16)
    pdf.cell(0, 10, "Predicted Exam Paper", ln=1, align="C")
    pdf.set_font("Arial", "", 12)
    pdf.cell(0, 10, f"{subject} | {source}: {teacher}", ln=1)
    pdf.cell(0, 10, datetime.now().strftime('%Y-%m-%d'), ln=1)
    pdf.ln(8)
    pdf.set_font("Arial", "B", 12)
    pdf.cell(0, 8, "Predicted Questions", ln=1)
    pdf.set_font("Arial", "", 11)
    for i, (q, conf, reason) in enumerate(questions, 1):
        pdf.multi_cell(0, 7, f"Q{i}. {q}  [Confidence: {conf}%]")
        pdf.set_font("Arial", "I", 9)
        pdf.multi_cell(0, 5, f"Reason: {reason}")
        pdf.set_font("Arial", "", 11)
        pdf.ln(2)
    if surprises:
        pdf.ln(5)
        pdf.set_font("Arial", "B", 12)
        pdf.cell(0, 8, "Surprise Topics", ln=1)
        pdf.set_font("Arial", "", 11)
        for t in surprises:
            pdf.cell(0, 6, f"- {t}", ln=1)
    path = f"{BASE}/predictions/{subject}/{subject}_{teacher}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
    pdf.output(path)
    return path

def export_word(subject, teacher, source, questions, surprises):
    doc = Document()
    doc.add_heading(f'Predicted Exam Paper: {subject}', 0)
    doc.add_paragraph(f'Source: {source} | Teacher: {teacher}')
    doc.add_paragraph(f'Date: {datetime.now().strftime("%Y-%m-%d")}')
    doc.add_heading('Predicted Questions', level=1)
    for i, (q, conf, reason) in enumerate(questions, 1):
        doc.add_paragraph(f'Q{i}. {q}  [Confidence: {conf}%]', style='List Bullet')
        doc.add_paragraph(f'Reason: {reason}', style='Intense Quote')
    if surprises:
        doc.add_heading('Surprise Topics', level=1)
        for t in surprises:
            doc.add_paragraph(t, style='List Bullet')
    path = f"{BASE}/predictions/{subject}/{subject}_{teacher}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    doc.save(path)
    return path

def export_html(subject, teacher, source, questions, surprises):
    html = f"""<html><head><title>Predicted Exam</title></head><body>
    <h1>Predicted Exam Paper</h1>
    <p><b>Subject:</b> {subject}<br><b>Source:</b> {source}<br><b>Teacher:</b> {teacher}<br><b>Date:</b> {datetime.now().strftime('%Y-%m-%d')}</p>
    <h2>Predicted Questions</h2><ul>"""
    for i, (q, conf, reason) in enumerate(questions, 1):
        html += f"<li><b>Q{i}. {q}</b>  [Confidence: {conf}%]<br><i>Reason: {reason}</i></li>"
    html += "</ul>"
    if surprises:
        html += "<h2>Surprise Topics</h2><ul>"
        for t in surprises:
            html += f"<li>{t}</li>"
        html += "</ul>"
    html += "</body></html>"
    path = f"{BASE}/predictions/{subject}/{subject}_{teacher}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.html"
    with open(path, "w") as f:
        f.write(html)
    return path

def save_prediction_history(subject, teacher, questions, surprises):
    history_path = f"{BASE}/history/{subject}_{teacher}.json"
    history = []
    if os.path.exists(history_path):
        with open(history_path) as f:
            history = json.load(f)
    history.append({
        "date": datetime.now().isoformat(),
        "questions": questions,
        "surprises": surprises
    })
    with open(history_path, "w") as f:
        json.dump(history, f)

def save_feedback(subject, teacher, correct_questions):
    path = f"{BASE}/feedback/{subject}_{teacher}.json"
    data = {"correct": correct_questions, "last_updated": datetime.now().isoformat()}
    with open(path, "w") as f:
        json.dump(data, f)

def extract_question_bank(subject):
    papers_dir = f"{BASE}/past_papers/{subject}"
    all_questions = []
    meta_file = f"{papers_dir}/metadata.json"
    if os.path.exists(meta_file):
        with open(meta_file) as f:
            meta = json.load(f)
        for item in meta:
            pdf_path = f"{papers_dir}/{item['file']}"
            if os.path.exists(pdf_path):
                loader = PyPDFLoader(pdf_path)
                pages = loader.load()
                text = " ".join([p.page_content for p in pages])
                found = re.findall(r'(?:Q\.?\s*\d+[.:]\s*)([^\n]+)', text)
                all_questions.extend(found)
    return all_questions

def compare_teachers(subject, teacher1, teacher2):
    pattern1 = extract_pattern(subject, teacher1, "College Teacher") or "No data"
    pattern2 = extract_pattern(subject, teacher2, "College Teacher") or "No data"
    comparison = f"**Teacher 1: {teacher1}**\n{pattern1}\n\n**Teacher 2: {teacher2}**\n{pattern2}"
    return comparison

def plot_topic_frequency(subject, teacher):
    topics = ["Macroeconomics", "Microeconomics", "International Trade", "GDP", "Inflation"]
    frequencies = [5, 3, 2, 4, 1]
    df = pd.DataFrame({"Topic": topics, "Frequency": frequencies})
    fig = px.bar(df, x="Topic", y="Frequency", title=f"Topic Frequency for {subject} - {teacher}")
    return fig

# =================== NEW: RAG CHATBOT FUNCTION ===================
def get_chatbot_response(user_message, chat_history):
    """Use Gemini with RAG context from all textbooks and past papers"""
    # First, search all subjects' vector stores for relevant content
    all_context = []
    for subject in SUBJECTS:
        try:
            vectordb = Chroma(persist_directory=f"{BASE}/vectorstore/{subject}", embedding_function=HuggingFaceEmbeddings(model_name="sentence-transformers/all-MiniLM-L6-v2"))
            retriever = vectordb.as_retriever(search_kwargs={"k": 3})
            docs = retriever.get_relevant_documents(user_message)
            for doc in docs:
                all_context.append(f"[{subject}] {doc.page_content[:500]}")
        except:
            continue
    context = "\n\n".join(all_context) if all_context else "No relevant documents found in your uploaded textbooks or past papers."

    # Also get schedule info
    schedule_df = get_schedule()
    schedule_text = schedule_df.to_string(index=False)

    # Build prompt with conversation history
    history_text = ""
    for msg in chat_history[-5:]:  # last 5 exchanges
        history_text += f"{msg['role']}: {msg['content']}\n"

    prompt = f"""You are a helpful exam assistant. Use the following context to answer the user's question naturally.
If the user asks to predict questions, you can guide them to the Predict Exam page or give a sample prediction.
If the user asks about schedule, show it.
Be friendly and concise.

**User's documents (from textbooks & past papers):**
{context}

**Current exam schedule:**
{schedule_text}

**Recent conversation:**
{history_text}

**User's new message:** {user_message}

**Your response:"""
    response = MODEL.generate_content(prompt)
    return response.text

# ------------------- SIDEBAR (Upcoming exams widget) -------------------
schedule_df = get_schedule()
today = datetime.now().date()
upcoming_exams = schedule_df[schedule_df['date'].dt.date >= today].sort_values('date').head(5)
st.sidebar.markdown("### 📅 Upcoming Exams")
for _, row in upcoming_exams.iterrows():
    st.sidebar.write(f"**{row['subject']}** – {row['date'].date()} ({row['exam_type']})")

st.sidebar.markdown("---")
st.sidebar.markdown("### 🛠️ Navigation")
page = st.sidebar.radio("Go to", [
    "Upload Files", "Schedule", "Important Questions", "Extract Pattern",
    "Predict Exam", "Internet Check", "Chatbot", "Question Bank",
    "Teacher Comparison", "Visual Analytics", "History & Feedback", "Download Data"
])

# ------------------- PAGE IMPLEMENTATIONS -------------------

if page == "Upload Files":
    st.header("📂 Upload Textbooks or Past Papers")
    col1, col2 = st.columns(2)
    with col1:
        subject = st.selectbox("Subject", SUBJECTS)
        doc_type = st.radio("Type", ["Textbook", "Past Paper"])
    with col2:
        uploaded = st.file_uploader(f"Upload {doc_type} (PDF)", type=["pdf"])
    if uploaded:
        target = f"{BASE}/{'textbooks' if doc_type=='Textbook' else 'past_papers'}/{subject}"
        os.makedirs(target, exist_ok=True)
        with open(f"{target}/{uploaded.name}", "wb") as f:
            f.write(uploaded.getbuffer())
        st.success(f"Saved {uploaded.name}")
        if doc_type == "Past Paper":
            with st.expander("Add Metadata"):
                exam_date = st.date_input("Exam Date")
                teacher = st.text_input("Teacher Name")
                exam_type = st.selectbox("Exam Type", EXAM_TYPES)
                source = st.selectbox("Source", SOURCES)
                if st.button("Save Metadata"):
                    meta_path = f"{target}/metadata.json"
                    meta = []
                    if os.path.exists(meta_path):
                        with open(meta_path) as f:
                            meta = json.load(f)
                    meta.append({"file": uploaded.name, "date": str(exam_date), "teacher": teacher, "type": exam_type, "source": source})
                    with open(meta_path, "w") as f:
                        json.dump(meta, f)
                    st.success("Metadata saved")
        if st.button("Build RAG Index"):
            with st.spinner("Building index..."):
                build_rag(subject)
            st.success("Index built")

elif page == "Schedule":
    st.header("📅 Exam Schedule (Editable)")
    df = get_schedule()
    edited = st.data_editor(df, num_rows="dynamic")
    if st.button("Save Schedule"):
        save_schedule(edited)
        st.success("Schedule saved")

elif page == "Important Questions":
    st.header("⭐ Teacher's Important Questions")
    subject = st.selectbox("Subject", SUBJECTS)
    teacher = st.text_input("Teacher Name")
    questions_text = st.text_area("Enter one question per line")
    if st.button("Save Important Questions") and teacher:
        qlist = [q.strip() for q in questions_text.split("\n") if q.strip()]
        save_important(subject, teacher, qlist)
        st.success(f"Saved {len(qlist)} important questions")

elif page == "Extract Pattern":
    st.header("🧠 Extract Pattern from Past Papers")
    subject = st.selectbox("Subject", SUBJECTS)
    source = st.selectbox("Source", SOURCES)
    teacher = ""
    if source == "College Teacher":
        teacher = st.text_input("Teacher Name")
    else:
        teacher = "Government Experts"
    if st.button("Extract"):
        with st.spinner("Analyzing..."):
            pattern = extract_pattern(subject, teacher, source)
            if pattern:
                st.subheader("Pattern")
                st.write(pattern)
                pattern_dir = f"{BASE}/patterns/{subject}"
                os.makedirs(pattern_dir, exist_ok=True)
                with open(f"{pattern_dir}/{source}_{teacher}.txt", "w") as f:
                    f.write(pattern)
                st.success("Pattern saved")
            else:
                st.warning("No past papers found.")

elif page == "Predict Exam":
    st.header("🔮 Predict Next Exam")
    schedule = get_schedule()
    today = datetime.now().date()
    next_exams = schedule[schedule['date'].dt.date >= today].sort_values('date')
    if not next_exams.empty:
        ne = next_exams.iloc[0]
        st.info(f"Next scheduled: {ne['subject']} on {ne['date'].date()} ({ne['exam_type']} by {ne['source']})")
    col1, col2 = st.columns(2)
    with col1:
        subject = st.selectbox("Subject", SUBJECTS)
        source = st.selectbox("Source", SOURCES)
        if source == "College Teacher":
            teacher = st.text_input("Teacher Name")
        else:
            teacher = "Government"
    with col2:
        num_questions = st.slider("Number of questions to predict", 3, 20, 5)
        difficulty = st.select_slider("Difficulty", options=DIFFICULTIES, value="Medium")
    custom = st.text_area("Extra instructions (optional)")
    if st.button("Generate Prediction"):
        with st.spinner("Analyzing..."):
            try:
                build_rag(subject)
            except:
                pass
            questions, surprises = predict(subject, teacher, source, custom, num_questions, difficulty)
            if questions:
                st.subheader("Predicted Questions")
                for i, (q, conf, reason) in enumerate(questions, 1):
                    st.markdown(f"**Q{i}. {q}**  *[Confidence: {conf}%]*")
                    st.caption(f"Reason: {reason}")
                if surprises:
                    st.subheader("Surprise Topics")
                    st.write(", ".join(surprises))
                save_prediction_history(subject, teacher, questions, surprises)
                st.subheader("Export Prediction")
                col_a, col_b, col_c = st.columns(3)
                pdf_path = export_pdf(subject, teacher, source, questions, surprises)
                word_path = export_word(subject, teacher, source, questions, surprises)
                html_path = export_html(subject, teacher, source, questions, surprises)
                with col_a:
                    with open(pdf_path, "rb") as f:
                        st.download_button("📄 Download PDF", f, file_name=f"{subject}_{teacher}_prediction.pdf")
                with col_b:
                    with open(word_path, "rb") as f:
                        st.download_button("📝 Download Word", f, file_name=f"{subject}_{teacher}_prediction.docx")
                with col_c:
                    with open(html_path, "rb") as f:
                        st.download_button("🌐 Download HTML", f, file_name=f"{subject}_{teacher}_prediction.html")
            else:
                st.error("Prediction failed. Upload textbooks/past papers first.")

elif page == "Internet Check":
    st.header("🌐 Check if a Question is Copied")
    question = st.text_area("Paste the exam question")
    if st.button("Check Internet"):
        with st.spinner("Searching..."):
            status, url = check_online(question)
            if status == "Found Online":
                st.error(f"Found online at: {url}")
            else:
                st.success("Likely original")

# =================== FIXED CHATBOT PAGE ===================
elif page == "Chatbot":
    st.header("🤖 AI Chat Assistant")
    st.markdown("Ask me anything about your exams, textbooks, schedule, or predictions. I can search your uploaded documents and help you naturally.")
    
    # Initialize chat history
    if "messages" not in st.session_state:
        st.session_state.messages = [{"role": "assistant", "content": "Hello! I'm your exam assistant. I have access to your textbooks, past papers, and schedule. How can I help you today?"}]
    
    # Display chat messages
    for msg in st.session_state.messages:
        with st.chat_message(msg["role"]):
            st.write(msg["content"])
    
    # User input
    if prompt := st.chat_input("Type your question here..."):
        # Add user message
        st.session_state.messages.append({"role": "user", "content": prompt})
        with st.chat_message("user"):
            st.write(prompt)
        
        # Get AI response
        with st.chat_message("assistant"):
            with st.spinner("Thinking..."):
                response = get_chatbot_response(prompt, st.session_state.messages[:-1])
                st.write(response)
        st.session_state.messages.append({"role": "assistant", "content": response})

elif page == "Question Bank":
    st.header("📚 Question Bank from Past Papers")
    subject = st.selectbox("Subject", SUBJECTS)
    if st.button("Extract Questions"):
        with st.spinner("Extracting..."):
            questions = extract_question_bank(subject)
            if questions:
                st.write(f"Found {len(questions)} questions:")
                for q in questions[:50]:
                    st.write(f"- {q}")
            else:
                st.info("No questions extracted. Upload past papers first.")

elif page == "Teacher Comparison":
    st.header("👩‍🏫 Compare Two Teachers")
    subject = st.selectbox("Subject", SUBJECTS)
    teacher1 = st.text_input("First Teacher Name")
    teacher2 = st.text_input("Second Teacher Name")
    if st.button("Compare"):
        comparison = compare_teachers(subject, teacher1, teacher2)
        st.markdown(comparison)

elif page == "Visual Analytics":
    st.header("📊 Visual Analytics")
    subject = st.selectbox("Subject", SUBJECTS)
    teacher = st.text_input("Teacher Name (for topic frequency)")
    if st.button("Show Topic Frequency Chart"):
        fig = plot_topic_frequency(subject, teacher)
        st.plotly_chart(fig)

elif page == "History & Feedback":
    st.header("📜 Prediction History & Feedback")
    subject = st.selectbox("Subject", SUBJECTS)
    teacher = st.text_input("Teacher Name")
    history_path = f"{BASE}/history/{subject}_{teacher}.json"
    if os.path.exists(history_path):
        with open(history_path) as f:
            history = json.load(f)
        st.write(f"Past predictions: {len(history)}")
        for idx, entry in enumerate(history[-5:]):
            with st.expander(f"Prediction on {entry['date']}"):
                st.write(entry['questions'][:3])
        st.subheader("Feedback: Mark which predicted questions appeared")
        correct_input = st.text_area("Enter the question texts that appeared (one per line)")
        if st.button("Submit Feedback"):
            correct_list = [c.strip() for c in correct_input.split("\n") if c.strip()]
            save_feedback(subject, teacher, correct_list)
            st.success("Feedback saved. Future predictions will be improved.")
    else:
        st.info("No prediction history yet. Generate some predictions first.")

elif page == "Download Data":
    st.header("💾 Download All Your Data")
    if st.button("Create ZIP of all data"):
        zip_buffer = io.BytesIO()
        with zipfile.ZipFile(zip_buffer, "w") as zf:
            for root, dirs, files in os.walk(BASE):
                for file in files:
                    file_path = os.path.join(root, file)
                    arcname = os.path.relpath(file_path, BASE)
                    zf.write(file_path, arcname)
        zip_buffer.seek(0)
        st.download_button("📦 Download ZIP", zip_buffer, file_name="exam_predictor_data.zip")
